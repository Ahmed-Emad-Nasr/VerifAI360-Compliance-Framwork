"""
evidence_processor.py
----------------------
Turns any uploaded evidence artifact (policy doc, config screenshot,
scan report, ...) into plain text so it can be sent to the AI analyzer.

Supported inputs:
    .txt, .log, .csv, .json         -> read directly
    .pdf                            -> pdfplumber text extraction (falls
                                        back to OCR per page if the PDF is
                                        a scanned image with no text layer)
    .docx                           -> python-docx paragraph/table extraction
    .png/.jpg/.jpeg/.bmp/.tiff      -> pytesseract OCR

Cross-platform Tesseract discovery
-----------------------------------
Tesseract is a native OS binary, not a Python package, so `pip install
pytesseract` alone is never enough — the engine itself has to be installed
separately and pytesseract has to be able to find it. `_configure_tesseract_path()`
below handles this on Windows, macOS, and Linux without any manual PATH
editing, in this order:

  1. TESSERACT_CMD env var (e.g. set in .env) — always wins if set.
  2. Already resolvable on PATH (covers the normal case: `apt install
     tesseract-ocr` on Linux, `brew install tesseract` on macOS when the
     app is launched from a shell that has brew's PATH loaded).
  3. A short list of common install locations per OS, for the cases PATH
     resolution misses (Windows installer never touches PATH by default;
     GUI-launched macOS apps often don't inherit the shell's PATH even
     after `brew install`).

If none of these find a binary, pytesseract's normal TesseractNotFoundError
is caught in extract_text() below and turned into an OS-specific,
actionable install message.
"""

import os
import shutil
import sys

TEXT_EXTENSIONS = {".txt", ".log", ".csv", ".json", ".md", ".conf", ".cfg", ".yaml", ".yml"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp"}

# Common install locations per OS that aren't always on PATH.
_CANDIDATE_PATHS = {
    "win32": [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ],
    "darwin": [
        "/opt/homebrew/bin/tesseract",   # Homebrew on Apple Silicon
        "/usr/local/bin/tesseract",      # Homebrew on Intel Macs
        "/opt/local/bin/tesseract",      # MacPorts
    ],
    "linux": [
        "/usr/bin/tesseract",
        "/usr/local/bin/tesseract",
        "/snap/bin/tesseract",
    ],
}

_INSTALL_INSTRUCTIONS = (
    "- Windows: download & run the installer from "
    "https://github.com/UB-Mannheim/tesseract/wiki (default install path is picked up "
    "automatically — just restart the app afterwards).\n"
    "- macOS: brew install tesseract\n"
    "- Linux (Debian/Ubuntu): sudo apt install tesseract-ocr\n"
    "- Linux (Fedora): sudo dnf install tesseract\n"
    "- Linux (Arch): sudo pacman -S tesseract\n"
    "If you installed it somewhere non-standard, set TESSERACT_CMD in your .env file "
    "to the full path of the tesseract binary."
)

_tesseract_configured = False


def _os_key() -> str:
    if sys.platform.startswith("win"):
        return "win32"
    if sys.platform == "darwin":
        return "darwin"
    return "linux"


def _configure_tesseract_path():
    """Make sure pytesseract knows where the tesseract binary lives. Runs once per process."""
    global _tesseract_configured
    if _tesseract_configured:
        return
    _tesseract_configured = True

    try:
        import pytesseract
    except ImportError:
        return

    env_path = os.environ.get("TESSERACT_CMD")
    if env_path:
        pytesseract.pytesseract.tesseract_cmd = env_path
        return

    if shutil.which("tesseract"):
        return  # already on PATH, nothing to do

    for candidate in _CANDIDATE_PATHS.get(_os_key(), []):
        if os.path.isfile(candidate):
            pytesseract.pytesseract.tesseract_cmd = candidate
            return


class EvidenceExtractionError(Exception):
    pass


class EvidenceSignatureError(Exception):
    """Raised when a file's actual content doesn't match what its extension claims —
    e.g. a script renamed to evidence.pdf. See README section 12, known-gap #5."""
    pass


# Magic-byte (file signature) checks, independent of the filename extension.
# Only formats with a reliable, short, fixed signature are checked here —
# plain-text formats (.txt/.log/.csv/.json/...) have no such signature by
# design, so they're intentionally not included and are trusted as-is.
def _sniff_bytes(filepath: str, n: int = 16) -> bytes:
    with open(filepath, "rb") as f:
        return f.read(n)


def _is_pdf(head: bytes) -> bool:
    return head.startswith(b"%PDF-")


def _is_docx(head: bytes) -> bool:
    # .docx is a ZIP container (PK\x03\x04 or the empty-archive variant PK\x05\x06).
    return head.startswith(b"PK\x03\x04") or head.startswith(b"PK\x05\x06")


def _is_png(head: bytes) -> bool:
    return head.startswith(b"\x89PNG\r\n\x1a\n")


def _is_jpeg(head: bytes) -> bool:
    return head.startswith(b"\xff\xd8\xff")


def _is_bmp(head: bytes) -> bool:
    return head.startswith(b"BM")


def _is_tiff(head: bytes) -> bool:
    return head.startswith(b"II*\x00") or head.startswith(b"MM\x00*")


def _is_webp(head: bytes) -> bool:
    return head[:4] == b"RIFF" and head[8:12] == b"WEBP"


_SIGNATURE_CHECKS = {
    ".pdf": (_is_pdf, "PDF"),
    ".docx": (_is_docx, "ZIP/DOCX"),
    ".png": (_is_png, "PNG"),
    ".jpg": (_is_jpeg, "JPEG"),
    ".jpeg": (_is_jpeg, "JPEG"),
    ".bmp": (_is_bmp, "BMP"),
    ".tiff": (_is_tiff, "TIFF"),
    ".tif": (_is_tiff, "TIFF"),
    ".webp": (_is_webp, "WEBP"),
}


def validate_file_signature(filepath: str, claimed_filename: str) -> None:
    """
    Content-based validation, not just trusting the extension: for every
    format with a reliable magic-byte signature (PDF, DOCX, and the common
    image types), read the first few bytes on disk and confirm they match
    what the extension claims. Raises EvidenceSignatureError on a mismatch
    (e.g. a renamed .exe uploaded as "screenshot.png") rather than silently
    handing whatever-it-actually-is to pdfplumber/python-docx/pytesseract.

    Text-like extensions (.txt/.log/.csv/.json/.md/...) have no fixed
    signature to check by design and are intentionally skipped here.
    """
    ext = os.path.splitext(claimed_filename)[1].lower()
    check = _SIGNATURE_CHECKS.get(ext)
    if not check:
        return  # no signature defined for this extension (e.g. plain text) — nothing to check
    matcher, expected_kind = check
    try:
        head = _sniff_bytes(filepath)
    except OSError as e:
        raise EvidenceSignatureError(f"Could not read file to verify its contents: {e}") from e
    if not matcher(head):
        raise EvidenceSignatureError(
            f"'{claimed_filename}' has a .{ext.lstrip('.')} extension, but its contents don't look "
            f"like a real {expected_kind} file. This upload was rejected — a renamed/mismatched file "
            f"can't be processed as {expected_kind}."
        )


def detect_evidence_type(filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    if ext in IMAGE_EXTENSIONS:
        return "screenshot/image"
    if ext == ".pdf":
        return "pdf_document"
    if ext == ".docx":
        return "word_document"
    if ext in TEXT_EXTENSIONS:
        return "text/log/config"
    return "unknown"


def extract_text(filepath: str) -> str:
    """Return best-effort plain text content of the evidence file."""
    ext = os.path.splitext(filepath)[1].lower()

    if ext in TEXT_EXTENSIONS:
        return _read_plain_text(filepath)
    if ext == ".pdf":
        return _extract_pdf(filepath)
    if ext == ".docx":
        return _extract_docx(filepath)
    if ext in IMAGE_EXTENSIONS:
        return _extract_image_ocr(filepath)

    raise EvidenceExtractionError(
        f"Unsupported evidence file type: '{ext}'. Supported: text/log/config, "
        f".pdf, .docx, and common image formats (screenshots)."
    )


def _read_plain_text(filepath: str) -> str:
    with open(filepath, "r", errors="ignore") as f:
        return f.read()


def _extract_pdf(filepath: str) -> str:
    try:
        import pdfplumber
    except ImportError as e:
        raise EvidenceExtractionError(
            "pdfplumber is not installed. Run: pip install pdfplumber"
        ) from e

    text_parts = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            if not page_text.strip():
                # Likely a scanned page with no text layer -> OCR it.
                page_text = _ocr_pdf_page(page)
            text_parts.append(page_text)
    return "\n".join(text_parts)


def _ocr_pdf_page(page) -> str:
    try:
        import pytesseract
        _configure_tesseract_path()
        im = page.to_image(resolution=200).original
        return pytesseract.image_to_string(im)
    except Exception:
        return ""


def _extract_docx(filepath: str) -> str:
    try:
        import docx
    except ImportError as e:
        raise EvidenceExtractionError(
            "python-docx is not installed. Run: pip install python-docx"
        ) from e

    document = docx.Document(filepath)
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _extract_image_ocr(filepath: str) -> str:
    try:
        import pytesseract
        from PIL import Image
    except ImportError as e:
        raise EvidenceExtractionError(
            "pytesseract/Pillow are not installed. Run: pip install pytesseract pillow "
            "(and install the tesseract-ocr system package)."
        ) from e

    _configure_tesseract_path()

    try:
        image = Image.open(filepath)
        text = pytesseract.image_to_string(image)
    except pytesseract.TesseractNotFoundError as e:
        raise EvidenceExtractionError(
            "Tesseract OCR engine is not installed (or not found) on this system.\n" + _INSTALL_INSTRUCTIONS
        ) from e
    return text
